# /// script
# requires-python = ">=3.12"
# dependencies = ["python-docx>=1.2,<2", "pillow>=11,<13", "click>=8.3,<9"]
# ///
"""Word document (.docx) creation and editing tool.

Commands: create, template, info, replace, extract.

Usage:
    uv run docx_tool.py COMMAND [OPTIONS]
"""

import json
import re
import sys
from pathlib import Path

import click
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


def write_output(text: str, output_path: str | None) -> None:
    """Write text to file or stdout."""
    if output_path:
        Path(output_path).write_text(text, encoding="utf-8")
        click.echo(f"Output written to {output_path}", err=True)
    else:
        click.echo(text)


def _load_json_input(value: str) -> object:
    candidate = Path(value)
    if candidate.exists() and candidate.is_file():
        return json.loads(candidate.read_text(encoding="utf-8"))
    return json.loads(value)


def _rgb(value: object) -> RGBColor | None:
    if value is None:
        return None
    text = str(value).strip().lstrip("#")
    if len(text) != 6:
        raise ValueError(f"invalid RGB color: {value}")
    return RGBColor.from_string(text.upper())


def _apply_run_spec(run, spec: dict[str, object]) -> None:
    if "bold" in spec:
        run.bold = bool(spec["bold"])
    if "italic" in spec:
        run.italic = bool(spec["italic"])
    if "underline" in spec:
        run.underline = bool(spec["underline"])
    if spec.get("fontName"):
        run.font.name = str(spec["fontName"])
    if spec.get("fontSize") is not None:
        run.font.size = Pt(float(spec["fontSize"]))
    if spec.get("color"):
        run.font.color.rgb = _rgb(spec["color"])


def _apply_paragraph_spec(paragraph, spec: dict[str, object]) -> None:
    alignment = spec.get("alignment")
    if alignment:
        mapping = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if str(alignment) not in mapping:
            raise ValueError(f"unsupported paragraph alignment: {alignment}")
        paragraph.alignment = mapping[str(alignment)]
    fmt = paragraph.paragraph_format
    if spec.get("spaceBefore") is not None:
        fmt.space_before = Pt(float(spec["spaceBefore"]))
    if spec.get("spaceAfter") is not None:
        fmt.space_after = Pt(float(spec["spaceAfter"]))
    if spec.get("lineSpacing") is not None:
        fmt.line_spacing = float(spec["lineSpacing"])
    if spec.get("keepWithNext") is not None:
        fmt.keep_with_next = bool(spec["keepWithNext"])

    runs = spec.get("runs")
    if isinstance(runs, list):
        paragraph.clear()
        for run_spec in runs:
            if not isinstance(run_spec, dict):
                raise ValueError("paragraph runs must be objects")
            run = paragraph.add_run(str(run_spec.get("text", "")))
            _apply_run_spec(run, run_spec)
    elif "text" in spec:
        paragraph.text = str(spec.get("text", ""))


def _shade_cell(cell, color: object) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), str(color).lstrip("#").upper())
    cell._tc.get_or_add_tcPr().append(shading)


def _set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _configure_header_footer(section, spec: dict[str, object]) -> None:
    for key, container in (("header", section.header), ("footer", section.footer)):
        value = spec.get(key)
        if value is None:
            continue
        container.is_linked_to_previous = False
        paragraph = container.paragraphs[0] if container.paragraphs else container.add_paragraph()
        if isinstance(value, dict):
            paragraph.text = str(value.get("text", ""))
            _apply_paragraph_spec(paragraph, value)
            if value.get("pageNumber"):
                if paragraph.text:
                    paragraph.add_run(" ")
                _set_page_number(paragraph)
        else:
            paragraph.text = str(value)


def _configure_section(section, spec: dict[str, object]) -> None:
    page = spec.get("page", spec)
    if not isinstance(page, dict):
        return
    size = str(page.get("size", "")).lower()
    if size == "a4":
        section.page_width, section.page_height = Mm(210), Mm(297)
    elif size == "letter":
        section.page_width, section.page_height = Inches(8.5), Inches(11)
    if str(page.get("orientation", "portrait")).lower() == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width
    margins = page.get("margins")
    if isinstance(margins, dict):
        for key, attr in (("top", "top_margin"), ("right", "right_margin"), ("bottom", "bottom_margin"), ("left", "left_margin")):
            if margins.get(key) is not None:
                setattr(section, attr, Inches(float(margins[key])))
    _configure_header_footer(section, spec)


def _apply_document_styles(doc: Document, styles: object) -> None:
    if not isinstance(styles, dict):
        return
    for style_name, style_spec in styles.items():
        if style_name not in doc.styles or not isinstance(style_spec, dict):
            continue
        style = doc.styles[style_name]
        if style_spec.get("fontName"):
            style.font.name = str(style_spec["fontName"])
        if style_spec.get("fontSize") is not None:
            style.font.size = Pt(float(style_spec["fontSize"]))
        if style_spec.get("color"):
            style.font.color.rgb = _rgb(style_spec["color"])
        if style_spec.get("bold") is not None:
            style.font.bold = bool(style_spec["bold"])


def _build_from_spec(doc: Document, spec: dict[str, object]) -> None:
    _apply_document_styles(doc, spec.get("styles"))
    if isinstance(spec.get("properties"), dict):
        for key in ("title", "subject", "author", "keywords", "comments", "category"):
            if key in spec["properties"]:
                setattr(doc.core_properties, key, str(spec["properties"][key]))
    _configure_section(doc.sections[0], spec)

    blocks = spec.get("blocks", [])
    if not isinstance(blocks, list):
        raise ValueError("document spec blocks must be an array")
    for block in blocks:
        if not isinstance(block, dict):
            raise ValueError("document blocks must be objects")
        block_type = str(block.get("type", "paragraph"))
        if block_type == "heading":
            paragraph = doc.add_heading(str(block.get("text", "")), level=int(block.get("level", 1)))
            _apply_paragraph_spec(paragraph, block)
        elif block_type in ("paragraph", "bullet", "number"):
            style = block.get("style")
            if not style and block_type == "bullet":
                style = "List Bullet"
            if not style and block_type == "number":
                style = "List Number"
            paragraph = doc.add_paragraph(style=str(style) if style else None)
            _apply_paragraph_spec(paragraph, block)
        elif block_type == "table":
            rows = block.get("rows")
            if not isinstance(rows, list) or not rows or any(not isinstance(row, list) for row in rows):
                raise ValueError("table rows must be a non-empty two-dimensional array")
            column_count = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=column_count)
            table.style = str(block.get("style", "Table Grid"))
            alignment = str(block.get("alignment", "left"))
            table.alignment = {"left": WD_TABLE_ALIGNMENT.LEFT, "center": WD_TABLE_ALIGNMENT.CENTER, "right": WD_TABLE_ALIGNMENT.RIGHT}.get(alignment, WD_TABLE_ALIGNMENT.LEFT)
            for row_index, row in enumerate(rows):
                for column_index, value in enumerate(row):
                    cell = table.cell(row_index, column_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    if isinstance(value, dict):
                        cell.text = str(value.get("text", ""))
                        if value.get("shading"):
                            _shade_cell(cell, value["shading"])
                        if isinstance(value.get("paragraph"), dict):
                            _apply_paragraph_spec(cell.paragraphs[0], value["paragraph"])
                    else:
                        cell.text = str(value)
            widths = block.get("columnWidths")
            if isinstance(widths, list):
                for row in table.rows:
                    for index, width in enumerate(widths[:len(row.cells)]):
                        row.cells[index].width = Inches(float(width))
        elif block_type == "image":
            if not block.get("path"):
                raise ValueError("image block requires path")
            paragraph = doc.add_paragraph()
            if block.get("alignment"):
                _apply_paragraph_spec(paragraph, {"alignment": block["alignment"]})
            run = paragraph.add_run()
            kwargs = {}
            if block.get("width") is not None:
                kwargs["width"] = Inches(float(block["width"]))
            if block.get("height") is not None:
                kwargs["height"] = Inches(float(block["height"]))
            run.add_picture(str(block["path"]), **kwargs)
            if block.get("caption"):
                caption = doc.add_paragraph(str(block["caption"]), style="Caption")
                caption.alignment = paragraph.alignment
        elif block_type == "pageBreak":
            doc.add_page_break()
        elif block_type == "sectionBreak":
            section_type = WD_SECTION.CONTINUOUS if block.get("continuous") else WD_SECTION.NEW_PAGE
            section = doc.add_section(section_type)
            _configure_section(section, block)
        else:
            raise ValueError(f"unsupported document block type: {block_type}")


def markdown_to_docx(md_text: str, doc: Document) -> None:
    """Convert simple markdown text to docx paragraphs.

    Supports: # headings (h1-h6), **bold**, *italic*, - bullet lists,
    1. numbered lists, blank lines as paragraph breaks, and --- as page breaks.
    """
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Page break
        if stripped in ("---", "***", "___"):
            doc.add_page_break()
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet_match:
            text = bullet_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            _apply_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if num_match:
            text = num_match.group(1)
            p = doc.add_paragraph(style="List Number")
            _apply_inline_formatting(p, text)
            i += 1
            continue

        # Empty line - skip
        if not stripped:
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _apply_inline_formatting(p, stripped)
        i += 1


def _apply_inline_formatting(paragraph, text: str) -> None:
    """Apply bold and italic inline formatting to a paragraph."""
    # Clear any default runs
    paragraph.clear()

    # Pattern to match **bold**, *italic*, ***bold italic***
    pattern = r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+|\*(?!\*)))"

    for match in re.finditer(pattern, text):
        if match.group(2):  # ***bold italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True
        elif match.group(5):  # plain text
            paragraph.add_run(match.group(5))


@click.group()
def cli() -> None:
    """Word document (.docx) creation and editing tool."""
    pass


@cli.command()
@click.option("--from-file", type=click.Path(exists=True, dir_okay=False), default=None, help="Input text or markdown file.")
@click.option("--text", type=str, default=None, help="Direct text content (supports basic markdown).")
@click.option("--json-data", type=str, default=None, help="Structured document JSON spec or path to a JSON spec file.")
@click.option("--title", type=str, default=None, help="Document title.")
@click.option("--font-size", type=float, default=11, help="Base font size in points (default: 11).")
@click.option("-o", "--output", type=click.Path(), required=True, help="Output .docx file path.")
def create(from_file: str | None, text: str | None, json_data: str | None, title: str | None, font_size: float, output: str) -> None:
    """Create a new Word document from text or markdown.

    Provide content via --from-file (reads a .txt/.md file) or --text (inline string).
    Basic markdown formatting is supported: headings, bold, italic, lists.
    """
    if sum(value is not None for value in (from_file, text, json_data)) != 1:
        click.echo("Error: provide exactly one of --from-file, --text, or --json-data.", err=True)
        sys.exit(1)

    try:
        doc = Document()

        if json_data:
            parsed = _load_json_input(json_data)
            if not isinstance(parsed, dict):
                raise ValueError("document JSON spec must be an object")
            _build_from_spec(doc, parsed)
        else:
            # Set default font size for markdown/text creation.
            style = doc.styles["Normal"]
            style.font.size = Pt(font_size)
            if title:
                doc.add_heading(title, level=0)
            if from_file:
                content = Path(from_file).read_text(encoding="utf-8")
            else:
                content = text or ""

            markdown_to_docx(content, doc)

        output_path = Path(output)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        click.echo(f"Document created: {output}", err=True)
    except (json.JSONDecodeError, ValueError) as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)
    except Exception as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)


@cli.command()
@click.argument("template_file", type=click.Path(exists=True, dir_okay=False))
@click.option("--data", type=str, required=True, help="JSON string or path to JSON file with template values.")
@click.option("-o", "--output", type=click.Path(), required=True, help="Output .docx file path.")
def template(template_file: str, data: str, output: str) -> None:
    """Fill a Word document template with JSON data.

    Replaces {{placeholder}} patterns in the document with values from JSON.
    The JSON should map placeholder names to values, e.g. {"name": "John", "date": "2024-01-01"}.
    """
    try:
        # Parse data
        data_path = Path(data)
        if data_path.exists() and data_path.is_file():
            template_data = json.loads(data_path.read_text(encoding="utf-8"))
        else:
            template_data = json.loads(data)

        if not isinstance(template_data, dict):
            click.echo("Error: JSON data must be an object.", err=True)
            sys.exit(1)

        doc = Document(template_file)

        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            _replace_in_paragraph(paragraph, template_data)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, template_data)

        # Replace in headers/footers
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                if not header_footer.is_linked_to_previous:
                    for paragraph in header_footer.paragraphs:
                        _replace_in_paragraph(paragraph, template_data)

        doc.save(output)
        click.echo(f"Template filled and saved to {output}", err=True)
    except json.JSONDecodeError as e:
        click.echo(f"Error parsing JSON: {e}", err=True)
        sys.exit(1)
    except Exception as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)


def _replace_in_paragraph(paragraph, data: dict[str, str]) -> None:
    """Replace {{key}} placeholders in a paragraph while preserving formatting."""
    full_text = paragraph.text
    if "{{" not in full_text:
        return

    # Single-pass replacement to avoid cascading (a replacement value containing
    # {{other_key}} should not be substituted again).
    def _replacer(match: re.Match) -> str:
        key = match.group(1)
        if key in data:
            return str(data[key])
        return match.group(0)  # leave unmatched placeholders as-is

    full_text = re.sub(r"\{\{(\w+)\}\}", _replacer, full_text)

    # Rebuild runs with the replaced text
    if paragraph.runs:
        # Clear all runs except the first, put full text in first run
        for i in range(len(paragraph.runs) - 1, 0, -1):
            paragraph.runs[i].text = ""
        paragraph.runs[0].text = full_text


@cli.command()
@click.argument("file", type=click.Path(exists=True, dir_okay=False))
@click.option("-o", "--output", type=click.Path(), default=None, help="Write output to file.")
def info(file: str, output: str | None) -> None:
    """Show document information: paragraphs, sections, tables, styles."""
    try:
        doc = Document(file)

        # Count elements
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        section_count = len(doc.sections)

        # Styles used
        styles_used = set()
        for p in doc.paragraphs:
            if p.style and p.style.name:
                styles_used.add(p.style.name)

        # Word count
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)

        # Core properties
        props = doc.core_properties
        info_dict: dict[str, object] = {
            "file": str(Path(file).resolve()),
            "paragraphs": paragraph_count,
            "tables": table_count,
            "sections": section_count,
            "word_count": word_count,
            "styles_used": sorted(styles_used),
            "inline_shapes": len(doc.inline_shapes),
            "section_details": [
                {
                    "width_inches": round(section.page_width / 914400, 2),
                    "height_inches": round(section.page_height / 914400, 2),
                    "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait",
                    "margins_inches": {
                        "top": round(section.top_margin / 914400, 2),
                        "right": round(section.right_margin / 914400, 2),
                        "bottom": round(section.bottom_margin / 914400, 2),
                        "left": round(section.left_margin / 914400, 2),
                    },
                    "header": "\n".join(p.text for p in section.header.paragraphs if p.text),
                    "footer": "\n".join(p.text for p in section.footer.paragraphs if p.text),
                }
                for section in doc.sections
            ],
            "table_details": [
                {"rows": len(table.rows), "columns": len(table.columns), "style": table.style.name if table.style else None}
                for table in doc.tables
            ],
            "properties": {
                "title": props.title,
                "author": props.author,
                "subject": props.subject,
                "created": str(props.created) if props.created else None,
                "modified": str(props.modified) if props.modified else None,
                "last_modified_by": props.last_modified_by,
                "revision": props.revision,
            },
        }

        # Template placeholders
        placeholders = set()
        for p in doc.paragraphs:
            for match in re.finditer(r"\{\{(\w+)\}\}", p.text):
                placeholders.add(match.group(1))
        if placeholders:
            info_dict["template_placeholders"] = sorted(placeholders)

        write_output(json.dumps(info_dict, indent=2, default=str), output)
    except Exception as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)


@cli.command()
@click.argument("file", type=click.Path(exists=True, dir_okay=False))
@click.option("--find", type=str, required=True, help="Text to find.")
@click.option("--replace-with", type=str, required=True, help="Replacement text.")
@click.option("--case-sensitive/--no-case-sensitive", default=True, help="Case-sensitive matching (default: yes).")
@click.option("-o", "--output", type=click.Path(), required=True, help="Output .docx file path.")
def replace(file: str, find: str, replace_with: str, case_sensitive: bool, output: str) -> None:
    """Find and replace text in a Word document."""
    try:
        doc = Document(file)
        count = 0

        for paragraph in doc.paragraphs:
            c = _find_replace_paragraph(paragraph, find, replace_with, case_sensitive)
            count += c

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        c = _find_replace_paragraph(paragraph, find, replace_with, case_sensitive)
                        count += c

        # Replace in headers/footers
        for section in doc.sections:
            for header_footer in [section.header, section.footer]:
                if not header_footer.is_linked_to_previous:
                    for paragraph in header_footer.paragraphs:
                        c = _find_replace_paragraph(paragraph, find, replace_with, case_sensitive)
                        count += c

        doc.save(output)
        click.echo(f"Replaced {count} occurrence(s). Saved to {output}", err=True)
    except Exception as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)


def _find_replace_paragraph(paragraph, find: str, replace_with: str, case_sensitive: bool) -> int:
    """Replace text in a paragraph, returns count of replacements."""
    full_text = paragraph.text
    check_text = full_text if case_sensitive else full_text.lower()
    find_text = find if case_sensitive else find.lower()

    if find_text not in check_text:
        return 0

    count = check_text.count(find_text)

    if case_sensitive:
        new_text = full_text.replace(find, replace_with)
    else:
        # Case-insensitive replace
        pattern = re.compile(re.escape(find), re.IGNORECASE)
        new_text = pattern.sub(replace_with, full_text)

    # Rebuild runs
    if paragraph.runs:
        for i in range(len(paragraph.runs) - 1, 0, -1):
            paragraph.runs[i].text = ""
        paragraph.runs[0].text = new_text

    return count


@cli.command()
@click.argument("file", type=click.Path(exists=True, dir_okay=False))
@click.option("--include-tables/--no-tables", default=True, help="Include table content (default: yes).")
@click.option("-o", "--output", type=click.Path(), default=None, help="Write output to file.")
def extract(file: str, include_tables: bool, output: str | None) -> None:
    """Extract all text content from a Word document."""
    try:
        doc = Document(file)
        parts: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                # Add heading markers
                if paragraph.style and paragraph.style.name.startswith("Heading"):
                    try:
                        level = int(paragraph.style.name.split()[-1])
                        text = "#" * level + " " + text
                    except (ValueError, IndexError):
                        pass
                parts.append(text)

        if include_tables:
            for t_idx, table in enumerate(doc.tables):
                parts.append(f"\n[Table {t_idx + 1}]")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append(" | ".join(cells))

        write_output("\n".join(parts), output)
    except Exception as e:
        click.echo(f"Error: {e}", err=True)
        sys.exit(1)


if __name__ == "__main__":
    cli()
